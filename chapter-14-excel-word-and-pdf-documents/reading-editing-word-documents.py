import os
import docx

cwd = f'{os.getcwd()}/chapter-14-excel-word-and-pdf-documents'

doc = docx.Document(f'{cwd}/demo.docx')
print(doc.paragraphs)
print(doc.paragraphs[0].text)
print(doc.paragraphs[1].text)

p = doc.paragraphs[1]
print(p.runs)
print(p.runs[0].text)
print(p.runs[1].text)
print(p.runs[2].text)

# None
print(p.runs[0].bold)
# True
print(p.runs[1].bold)
# True
print(p.runs[3].italic)
p.runs[3].underline = True
p.runs[3].text = 'italic and underline.'
doc.save(f'{cwd}/demo2.docx')

p.style = 'Title'
doc.save(f'{cwd}/demo3.docx')

doc4 = docx.Document()
doc4.add_paragraph('Hello this is a paragraph.')
doc4.add_paragraph('This is another paragraph.')
doc4.save(f'{cwd}/demo4.docx')

p = doc4.paragraphs[0]
p.add_run('This is a new run')
p.runs[1].bold = True
doc4.save(f'{cwd}/demo5.docx')

# ====================================

def getText(filename):
	doc = docx.Document(filename)
	fullText = []
	
	for paragraph in doc.paragraphs:
		fullText.append(paragraph.text)

	return '\n'.join(fullText)

print(getText(f'{cwd}/demo.docx'))